#!/usr/bin/env python3
"""Build reproducible sample thesis Word documents for template validation."""

from __future__ import annotations

import argparse
import shutil
import subprocess
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as docx_qn
    from docx.shared import Mm, Pt, RGBColor
    from PIL import Image, ImageDraw
except ImportError as exc:  # pragma: no cover - dependency guidance path
    raise SystemExit(
        "Missing dependency. Install with `python -m pip install python-docx lxml Pillow pypdf`."
    ) from exc

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from docx_utils import (
    LAYOUT,
)
from postprocess_docx import postprocess_docx
from check_format import check_docx


ROOT = SCRIPT_DIR.parent
TEMPLATE_DIR = ROOT / "template"
SOURCE_DIR = ROOT / "source"
FIGURE_DIR = ROOT / "assets" / "figures"
BUILD_DIR = ROOT / "generated"
RAW_DOCX = BUILD_DIR / "sample_raw.docx"
FINAL_DOCX = BUILD_DIR / "sample_final.docx"
REFERENCE_DOCX = TEMPLATE_DIR / "reference.docx"
COVER_TEMPLATE = TEMPLATE_DIR / "cover_template.docx"


def _ensure_dirs() -> None:
    for path in [TEMPLATE_DIR, SOURCE_DIR, FIGURE_DIR, BUILD_DIR, BUILD_DIR / "preview"]:
        path.mkdir(parents=True, exist_ok=True)
    for name in ["template_0.docx", "template_1.docx"]:
        source = ROOT / name
        target = TEMPLATE_DIR / name
        if source.exists() and not target.exists():
            shutil.copy2(source, target)


def _set_run_font(run, size_pt: float = 13.0, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(docx_qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic


def _set_style_color_black(style) -> None:
    rpr = style._element.get_or_add_rPr()
    color = rpr.find(docx_qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        rpr.append(color)
    color.attrib.clear()
    color.set(docx_qn("w:val"), "000000")


def _ensure_paragraph_color_black(paragraph) -> None:
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)


def _set_paragraph_format(paragraph, *, style_name: str | None = None, body: bool = False) -> None:
    if style_name:
        paragraph.style = style_name
    fmt = paragraph.paragraph_format
    if body:
        fmt.first_line_indent = Mm(LAYOUT.left_margin_mm - 17.5)  # 12.5 mm
        fmt.line_spacing = 1.25
        fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)


def _configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(docx_qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    _set_style_color_black(normal)
    normal.paragraph_format.first_line_indent = Mm(12.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    body_text = styles["Body Text"] if "Body Text" in styles else styles.add_style("Body Text", WD_STYLE_TYPE.PARAGRAPH)
    body_text.base_style = normal
    body_text.font.name = "Times New Roman"
    body_text._element.rPr.rFonts.set(docx_qn("w:eastAsia"), "Times New Roman")
    body_text.font.size = Pt(13)
    body_text.font.color.rgb = RGBColor(0, 0, 0)
    _set_style_color_black(body_text)
    body_text.paragraph_format.first_line_indent = Mm(12.5)
    body_text.paragraph_format.line_spacing = 1.25
    body_text.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    h1 = styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1._element.rPr.rFonts.set(docx_qn("w:eastAsia"), "Times New Roman")
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.underline = False
    h1.font.all_caps = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    _set_style_color_black(h1)
    h1.paragraph_format.page_break_before = True
    h1.paragraph_format.first_line_indent = Mm(12.5)
    h1.paragraph_format.line_spacing = 1.25
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(13)

    h2 = styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2._element.rPr.rFonts.set(docx_qn("w:eastAsia"), "Times New Roman")
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.underline = False
    h2.font.color.rgb = RGBColor(0, 0, 0)
    _set_style_color_black(h2)
    h2.paragraph_format.first_line_indent = Mm(12.5)
    h2.paragraph_format.line_spacing = 1.25
    h2.paragraph_format.space_before = Pt(13)
    h2.paragraph_format.space_after = Pt(13)

    h3 = styles["Heading 3"]
    h3.font.name = "Times New Roman"
    h3._element.rPr.rFonts.set(docx_qn("w:eastAsia"), "Times New Roman")
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.underline = False
    h3.font.color.rgb = RGBColor(0, 0, 0)
    _set_style_color_black(h3)
    h3.paragraph_format.first_line_indent = Mm(12.5)
    h3.paragraph_format.line_spacing = 1.25

    caption = styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(docx_qn("w:eastAsia"), "Times New Roman")
    caption.font.size = Pt(13)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor(0, 0, 0)
    _set_style_color_black(caption)
    caption.paragraph_format.first_line_indent = Mm(0)
    caption.paragraph_format.line_spacing = 1.25
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(13)

    for style_name, alignment in [
        ("Figure Caption", WD_ALIGN_PARAGRAPH.CENTER),
        ("Table Caption", WD_ALIGN_PARAGRAPH.LEFT),
        ("Formula", WD_ALIGN_PARAGRAPH.CENTER),
        ("Contents Title", WD_ALIGN_PARAGRAPH.CENTER),
        ("Contents Entry", WD_ALIGN_PARAGRAPH.LEFT),
    ]:
        style = styles[style_name] if style_name in styles else styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(docx_qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14 if style_name == "Contents Title" else 13)
        style.font.bold = style_name == "Contents Title"
        style.font.underline = False
        style.font.color.rgb = RGBColor(0, 0, 0)
        _set_style_color_black(style)
        style.paragraph_format.first_line_indent = Mm(0)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.alignment = alignment
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(13 if style_name in {"Figure Caption", "Formula"} else 0)
        if style_name == "Contents Entry":
            style.paragraph_format.tab_stops.add_tab_stop(
                Mm(165),
                alignment=WD_TAB_ALIGNMENT.RIGHT,
                leader=WD_TAB_LEADER.DOTS,
            )
        if style_name == "Formula":
            style.paragraph_format.tab_stops.add_tab_stop(
                Mm(165),
                alignment=WD_TAB_ALIGNMENT.RIGHT,
                leader=WD_TAB_LEADER.SPACES,
            )


def _set_doc_layout(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Mm(LAYOUT.page_width_mm)
        section.page_height = Mm(LAYOUT.page_height_mm)
        section.left_margin = Mm(LAYOUT.left_margin_mm)
        section.right_margin = Mm(LAYOUT.right_margin_mm)
        section.top_margin = Mm(LAYOUT.top_margin_mm)
        section.bottom_margin = Mm(LAYOUT.bottom_margin_mm)
        section.footer_distance = Mm(LAYOUT.footer_distance_mm)
        section.header_distance = Mm(LAYOUT.header_distance_mm)


def build_reference_docx(path: Path = REFERENCE_DOCX) -> None:
    doc = Document()
    _set_doc_layout(doc)
    _configure_styles(doc)
    doc.add_paragraph("Reference style sample paragraph.", style="Body Text")
    doc.add_paragraph("1 SAMPLE SECTION", style="Heading 1")
    doc.add_paragraph("1.1 Sample subsection", style="Heading 2")
    doc.add_paragraph("1.1.1 Sample paragraph", style="Heading 3")
    doc.add_paragraph("Figure 1.1 – Sample caption", style="Figure Caption")
    doc.add_paragraph("Table 1.1 – Sample caption", style="Table Caption")
    add_omml_equation_with_number(doc, _heat_equation_omml(), "(1.1)")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def build_cover_template(path: Path = COVER_TEMPLATE) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _set_doc_layout(doc)
    _configure_styles(doc)
    section = doc.sections[0]
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    for paragraph in section.header.paragraphs + section.footer.paragraphs:
        paragraph.clear()
    if doc.paragraphs:
        doc.paragraphs[0].text = ""
    doc.save(path)
    path.chmod(0o644)


def _add_body_paragraph(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="Body Text")
    paragraph.paragraph_format.first_line_indent = Mm(12.5)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text)
    _set_run_font(run, 13)
    return paragraph


def _add_caption(doc: Document, text: str, style: str):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.first_line_indent = Mm(0)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if style == "Figure Caption" else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    _set_run_font(run, 13)
    return paragraph


def _make_sample_figure(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1100, 520), "white")
    draw = ImageDraw.Draw(img)
    draw.rectangle((40, 40, 1060, 480), outline=(30, 30, 30), width=4)
    boxes = [
        (100, 180, 320, 320, "Sensor"),
        (440, 180, 660, 320, "Edge hub"),
        (780, 180, 1000, 320, "Actuator"),
    ]
    for x1, y1, x2, y2, label in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=12, outline=(20, 20, 20), width=4, fill=(240, 246, 250))
        bbox = draw.textbbox((0, 0), label)
        draw.text(((x1 + x2 - bbox[2]) / 2, (y1 + y2 - bbox[3]) / 2), label, fill=(0, 0, 0))
    draw.line((320, 250, 440, 250), fill=(0, 0, 0), width=4)
    draw.line((660, 250, 780, 250), fill=(0, 0, 0), width=4)
    draw.polygon([(430, 240), (450, 250), (430, 260)], fill=(0, 0, 0))
    draw.polygon([(770, 240), (790, 250), (770, 260)], fill=(0, 0, 0))
    img.save(path)


def _add_contents_entry(doc: Document, title: str, page: str) -> None:
    paragraph = doc.add_paragraph(style="Contents Entry")
    paragraph.paragraph_format.first_line_indent = Mm(0)
    paragraph.paragraph_format.left_indent = Mm(0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Mm(165), alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.DOTS
    )
    run = paragraph.add_run(title)
    _set_run_font(run, 13)
    paragraph.add_run("\t")
    page_run = paragraph.add_run(page)
    _set_run_font(page_run, 13)


def _set_table_width(table, width_mm: float) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(docx_qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(docx_qn("w:type"), "dxa")
    tbl_w.set(docx_qn("w:w"), str(int(round(width_mm * 1440 / 25.4))))


def _omml_run(text: str) -> OxmlElement:
    run = OxmlElement("m:r")
    rpr = OxmlElement("m:rPr")
    run.append(rpr)
    text_el = OxmlElement("m:t")
    text_el.text = text
    run.append(text_el)
    return run


def _heat_equation_omml() -> str:
    return (
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        "<m:r><m:t>Q = mc</m:t></m:r>"
        "<m:r><m:t>ΔT</m:t></m:r>"
        "</m:oMath>"
    )


def add_omml_equation_with_number(doc: Document, equation_omml: str, number: str):
    paragraph = doc.add_paragraph(style="Formula")
    paragraph.paragraph_format.first_line_indent = Mm(0)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Mm(82.5), alignment=WD_TAB_ALIGNMENT.CENTER, leader=WD_TAB_LEADER.SPACES
    )
    ppr = paragraph._p.get_or_add_pPr()
    tabs = ppr.find(docx_qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        ppr.append(tabs)
    right_tab = OxmlElement("w:tab")
    right_tab.set(docx_qn("w:val"), "right")
    right_tab.set(docx_qn("w:leader"), "none")
    right_tab.set(docx_qn("w:pos"), str(int(round(165 * 1440 / 25.4))))
    tabs.append(right_tab)
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Mm(165), alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.SPACES
    )
    center_tab_run = paragraph.add_run("\t")
    _set_run_font(center_tab_run, 13)
    equation = parse_xml_fragment(equation_omml)
    paragraph._p.append(equation)
    tab_run = paragraph.add_run("\t")
    _set_run_font(tab_run, 13)
    number_run = paragraph.add_run(number)
    _set_run_font(number_run, 13)
    _ensure_paragraph_color_black(paragraph)
    return paragraph


def parse_xml_fragment(xml: str):
    from docx.oxml import parse_xml

    return parse_xml(xml)


def build_raw_docx(path: Path = RAW_DOCX) -> None:
    figure_path = FIGURE_DIR / "sample_system_diagram.png"
    _make_sample_figure(figure_path)

    doc = Document()
    _set_doc_layout(doc)
    _configure_styles(doc)

    # Cover section: intentionally blank; postprocess supplies the border-only template.
    doc.add_paragraph()

    doc.add_section(WD_SECTION.NEW_PAGE)

    contents_title = doc.add_paragraph(style="Contents Title")
    contents_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contents_title.add_run("Contents")
    _set_run_font(run, 14, bold=True)
    doc.add_paragraph()
    _add_contents_entry(doc, "1 SAMPLE TEMPLATE SECTION", "3")
    _add_contents_entry(doc, "1.1 Sample subsection", "3")
    _add_contents_entry(doc, "1.1.1 Sample paragraph", "4")

    doc.add_section(WD_SECTION.NEW_PAGE)

    heading = doc.add_paragraph(style="Heading 1")
    heading.add_run("1 SAMPLE TEMPLATE SECTION")
    for run in heading.runs:
        _set_run_font(run, 14, bold=True)

    heading2 = doc.add_paragraph(style="Heading 2")
    heading2.add_run("1.1 Sample subsection")
    for run in heading2.runs:
        _set_run_font(run, 13, bold=True)

    _add_body_paragraph(
        doc,
        "This sample paragraph exists only to validate the Word template system. "
        "In accordance with Figure 1.1, the generated document keeps figures centered, "
        "uses school margins, and stores page numbers as Word fields rather than plain text.",
    )
    _add_body_paragraph(
        doc,
        "The document also references Table 1.1 and formula (1.1) so the checker can "
        "confirm that captions and formula references are discoverable before the real thesis is written.",
    )

    image_paragraph = doc.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.first_line_indent = Mm(0)
    image_paragraph.add_run().add_picture(str(figure_path), width=Mm(120))
    _add_caption(doc, "Figure 1.1 – Structural diagram of the sample system", "Figure Caption")

    _add_caption(doc, "Table 1.1 – Sample validation data", "Table Caption")
    table = doc.add_table(rows=3, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    _set_table_width(table, 165)
    widths = [Mm(55), Mm(55), Mm(55)]
    headers = ["Parameter", "Value", "Note"]
    rows = [
        ["Page number", "Automatic field", "Not manually typed"],
        ["Missing data marker", "-", "Dash used"],
    ]
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.width = widths[col_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text = headers[col_idx] if row_idx == 0 else rows[row_idx - 1][col_idx]
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.first_line_indent = Mm(0)
            paragraph.paragraph_format.line_spacing = 1.25
            run = paragraph.add_run(text)
            _set_run_font(run, 13, bold=row_idx == 0)

    _add_body_paragraph(
        doc,
        "The sample equation is placed on its own line as a real Word math object and is "
        "referenced as formula (1.1) for automated validation.",
    )
    add_omml_equation_with_number(doc, _heat_equation_omml(), "(1.1)")
    for text in [
        "where Q – amount of heat, J;",
        "m – mass, kg;",
        "c – specific heat capacity, J/(kg·K);",
        "ΔT – temperature difference, K.",
    ]:
        _add_body_paragraph(doc, text)

    heading3 = doc.add_paragraph(style="Heading 3")
    heading3.add_run("1.1.1 Sample paragraph")
    for run in heading3.runs:
        _set_run_font(run, 13, bold=True)

    filler = (
        "This controlled filler paragraph extends the document so the generated body occupies "
        "multiple pages. The purpose is to verify that the right lower page marker follows "
        "Word pagination and is not a manually typed number. "
    )
    for idx in range(24):
        _add_body_paragraph(doc, f"{filler}Validation paragraph {idx + 1}.")

    for section in doc.sections:
        section.page_width = Mm(LAYOUT.page_width_mm)
        section.page_height = Mm(LAYOUT.page_height_mm)
        section.left_margin = Mm(LAYOUT.left_margin_mm)
        section.right_margin = Mm(LAYOUT.right_margin_mm)
        section.top_margin = Mm(LAYOUT.top_margin_mm)
        section.bottom_margin = Mm(LAYOUT.bottom_margin_mm)
        section.footer_distance = Mm(LAYOUT.footer_distance_mm)
        section.header_distance = Mm(LAYOUT.header_distance_mm)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def _write_source_stubs() -> None:
    files = {
        SOURCE_DIR / "00_cover.md": "# Cover\n\nThis placeholder is used by the template build only.\n",
        SOURCE_DIR / "01_contents.md": "# Contents\n\nGenerated automatically in the template sample.\n",
        SOURCE_DIR / "02_sample_body.md": (
            "# 1 SAMPLE TEMPLATE SECTION\n\n"
            "This sample file is not thesis prose. It documents the minimal body content "
            "used to validate the template system.\n"
        ),
        SOURCE_DIR / "metadata.yaml": (
            "title: Sample thesis template validation document\n"
            "author: Your_name\n"
            "department: Computer&Systems Department\n"
            "language: en\n"
        ),
    }
    for path, content in files.items():
        if not path.exists():
            path.write_text(content, encoding="utf-8")


def _try_render_preview(docx_path: Path) -> None:
    soffice = shutil.which("libreoffice") or shutil.which("soffice")
    if not soffice:
        return
    preview_dir = BUILD_DIR / "preview"
    preview_dir.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(preview_dir), str(docx_path)],
        check=False,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )


def build_all(strict: bool = True) -> int:
    _ensure_dirs()
    _write_source_stubs()
    build_reference_docx(REFERENCE_DOCX)
    build_cover_template(COVER_TEMPLATE)
    build_raw_docx(RAW_DOCX)
    postprocess_docx(RAW_DOCX, FINAL_DOCX, TEMPLATE_DIR)
    _try_render_preview(FINAL_DOCX)
    reporter = check_docx(FINAL_DOCX, BUILD_DIR / "format_report.md")
    counts = reporter.counts()
    print(f"Wrote {RAW_DOCX}")
    print(f"Wrote {FINAL_DOCX}")
    print(f"Wrote {BUILD_DIR / 'format_report.md'}")
    print(f"PASS: {counts['PASS']}")
    print(f"WARNING: {counts['WARNING']}")
    print(f"ERROR: {counts['ERROR']}")
    return 1 if strict and reporter.has_errors else 0


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--no-strict", action="store_true", help="Do not fail when format checker reports ERROR")
    args = parser.parse_args(argv)
    return build_all(strict=not args.no_strict)


if __name__ == "__main__":
    raise SystemExit(main())
