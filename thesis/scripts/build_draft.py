#!/usr/bin/env python3
"""Build the working thesis draft from approved Markdown source files."""

from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.oxml import OxmlElement
    from docx.oxml import parse_xml
    from docx.oxml.ns import qn as docx_qn
    from docx.shared import Mm, Pt, RGBColor
except ImportError as exc:  # pragma: no cover - dependency guidance path
    raise SystemExit("Missing dependency. Install with `python -m pip install python-docx lxml Pillow pypdf`.") from exc

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from check_format import check_docx
from docx_utils import LAYOUT
from postprocess_docx import postprocess_docx


ROOT = SCRIPT_DIR.parent
SOURCE_DIR = ROOT / "source" / "draft"
TEMPLATE_DIR = ROOT / "template"
DRAFT_DIR = ROOT / "generated" / "drafts"
RAW_DOCX = DRAFT_DIR / "thesis_draft_raw.docx"
FINAL_DOCX = DRAFT_DIR / "thesis_draft.docx"
REPORT_DOCX = DRAFT_DIR / "thesis_draft_format_report.md"
SAMPLE_TEMPLATE = ROOT / "generated" / "sample_final.docx"
FIGURE_3_1_SOURCE_PATH = ROOT.parent / "docs" / "figures" / "figure_3_1_general_architecture.png"
FIGURE_3_1_CROPPED_PATH = ROOT.parent / "docs" / "figures" / "figure_3_1_general_architecture_cropped.png"
FIGURE_3_1_PATH = FIGURE_3_1_CROPPED_PATH if FIGURE_3_1_CROPPED_PATH.exists() else FIGURE_3_1_SOURCE_PATH
FIGURES_DIR = ROOT.parent / "docs" / "figures"
FIGURE_DEFINITIONS = {
    "Figure 2.1": (
        FIGURE_3_1_PATH,
        "Figure 2.1 – General architecture of the developed system",
        165,
    ),
    "Figure 3.1": (
        FIGURES_DIR / "figure_4_1_schematic.png",
        "Figure 3.1 – Electrical schematic of the edge temperature-control node",
        150,
    ),
    "Figure 3.2": (
        FIGURES_DIR / "figure_4_2_pcb_design.png",
        "Figure 3.2 – Component placement of the edge control node",
        145,
    ),
    "Figure 3.3": (
        FIGURES_DIR / "figure_4_3_enclosure_layout.png",
        "Figure 3.3 – Three-dimensional enclosure layout with PCB reference",
        112,
    ),
    "Figure 3.4": (
        FIGURES_DIR / "figure_4_4_enclosure_parts.png",
        "Figure 3.4 – Printable enclosure parts prepared for later fabrication",
        135,
    ),
    "Figure 3.5": (
        FIGURES_DIR / "figure_4_5_pid_anti_windup_code.png",
        "Figure 3.5 – PID output limiting and anti-windup implementation fragment",
        165,
    ),
    "Figure 3.6": (
        FIGURES_DIR / "figure_4_6_edge_control_tick_code.png",
        "Figure 3.6 – Edge control tick with safety forcing and telemetry publishing",
        165,
    ),
    "Figure 4.1": (
        FIGURES_DIR / "figure_5_3_datahub_pipeline_code.png",
        "Figure 4.1 – Data Hub bounded MQTT ingestion implementation fragment",
        165,
    ),
    "Figure 4.2": (
        FIGURES_DIR / "figure_5_1_hmi_device_detail.png",
        "Figure 4.2 – HMI device detail screen for monitoring and parameter configuration",
        165,
    ),
    "Figure 4.3": (
        FIGURES_DIR / "figure_5_2_hmi_ops_console.png",
        "Figure 4.3 – HMI operations console for Data Hub and ingestion monitoring",
        165,
    ),
    "Figure 4.4": (
        FIGURES_DIR / "figure_5_4_hmi_command_publisher_code.png",
        "Figure 4.4 – HMI backend parameter command publishing fragment",
        165,
    ),
    "Figure 5.1": (
        FIGURES_DIR / "figure_6_2_feature_extraction_code.png",
        "Figure 5.1 – Feature extraction for control-behavior analysis",
        165,
    ),
    "Figure 5.2": (
        FIGURES_DIR / "figure_6_3_recommendation_rules_code.png",
        "Figure 5.2 – Rule-based parameter recommendation fragment",
        165,
    ),
    "Figure 5.3": (
        FIGURES_DIR / "figure_6_1_hmi_validation.png",
        "Figure 5.3 – HMI post-apply validation view with telemetry comparison",
        165,
    ),
    "Figure 5.4": (
        FIGURES_DIR / "figure_6_4_post_effect_evaluator_code.png",
        "Figure 5.4 – Post-apply effect evaluation fragment",
        165,
    ),
}
FORMULA_SYMBOL_TAB_MM = 16.0
CONTENTS_FIRST_PAGE_ENTRY_LIMIT = 27
NBSP = "\u00a0"
MANUAL_TABLE_SPLITS = {
    "Table 2.2": 2,
    "Table 3.1": 2,
    "Table 4.1": 2,
    "Table 5.2": 3,
}


def _keep_numbered_reference_together(text: str) -> str:
    text = re.sub(r"\b(Figure|Table) (?=\d+\.\d+\b)", rf"\1{NBSP}", text)
    return re.sub(
        r"\b(formula|equation|expression|equality|transfer function) (?=\(\d+\.\d+\))",
        lambda match: f"{match.group(1)}{NBSP}",
        text,
        flags=re.I,
    )


def _set_style_color_black(style) -> None:
    rpr = style._element.get_or_add_rPr()
    color = rpr.find(docx_qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        rpr.append(color)
    color.attrib.clear()
    color.set(docx_qn("w:val"), "000000")


def _set_run_font(run, size_pt: float = 13.0, *, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(docx_qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic


def _configure_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(docx_qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    _set_style_color_black(normal)
    normal.paragraph_format.first_line_indent = Mm(12.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    body_text = styles["Body Text"] if "Body Text" in styles else styles.add_style("Body Text", WD_STYLE_TYPE.PARAGRAPH)
    body_text.base_style = normal
    body_text.font.name = "Times New Roman"
    body_text._element.rPr.rFonts.set(docx_qn("w:eastAsia"), "Times New Roman")
    body_text.font.size = Pt(13)
    body_text.font.color.rgb = RGBColor(0, 0, 0)
    _set_style_color_black(body_text)
    body_text.paragraph_format.first_line_indent = Mm(12.5)
    body_text.paragraph_format.line_spacing = 1.25
    body_text.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body_text.paragraph_format.space_before = Pt(0)
    body_text.paragraph_format.space_after = Pt(0)

    h1 = styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1._element.rPr.rFonts.set(docx_qn("w:eastAsia"), "Times New Roman")
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.underline = False
    h1.font.all_caps = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    _set_style_color_black(h1)
    h1.paragraph_format.page_break_before = True
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.first_line_indent = Mm(12.5)
    h1.paragraph_format.line_spacing = 1.25
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(13)

    h2 = styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2._element.rPr.rFonts.set(docx_qn("w:eastAsia"), "Times New Roman")
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.underline = False
    h2.font.color.rgb = RGBColor(0, 0, 0)
    _set_style_color_black(h2)
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.first_line_indent = Mm(12.5)
    h2.paragraph_format.line_spacing = 1.25
    h2.paragraph_format.space_before = Pt(13)
    h2.paragraph_format.space_after = Pt(13)

    for style_name, size_pt, bold, alignment in [
        ("Contents Title", 14, True, WD_ALIGN_PARAGRAPH.CENTER),
        ("Contents Entry", 13, False, WD_ALIGN_PARAGRAPH.LEFT),
        ("Figure Caption", 13, False, WD_ALIGN_PARAGRAPH.CENTER),
        ("Table Caption", 13, False, WD_ALIGN_PARAGRAPH.LEFT),
        ("Formula", 13, False, WD_ALIGN_PARAGRAPH.LEFT),
        ("Caption", 13, False, WD_ALIGN_PARAGRAPH.LEFT),
        ("References Entry", 13, False, WD_ALIGN_PARAGRAPH.LEFT),
    ]:
        style = styles[style_name] if style_name in styles else styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(docx_qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size_pt)
        style.font.bold = bold
        style.font.underline = False
        style.font.color.rgb = RGBColor(0, 0, 0)
        _set_style_color_black(style)
        style.paragraph_format.first_line_indent = Mm(0)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.alignment = alignment
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        if style_name == "Contents Entry":
            style.paragraph_format.tab_stops.add_tab_stop(
                Mm(165), alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.DOTS
            )
        if style_name == "References Entry":
            style.paragraph_format.first_line_indent = Mm(0)
            style.paragraph_format.left_indent = Mm(0)


def _set_doc_layout(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Mm(LAYOUT.page_width_mm)
        section.page_height = Mm(LAYOUT.page_height_mm)
        section.left_margin = Mm(LAYOUT.left_margin_mm)
        section.right_margin = Mm(LAYOUT.right_margin_mm)
        section.top_margin = Mm(LAYOUT.top_margin_mm)
        section.bottom_margin = Mm(LAYOUT.bottom_margin_mm)
        section.footer_distance = Mm(LAYOUT.footer_distance_mm)
        section.header_distance = Mm(LAYOUT.header_distance_mm)


def _read_markdown_source() -> str:
    sources = sorted(SOURCE_DIR.glob("[0-9][0-9]_*.md"))
    return "\n\n".join(path.read_text(encoding="utf-8") for path in sources if path.exists())


def _extract_headings(markdown: str) -> list[tuple[int, str]]:
    headings = []
    for line in markdown.splitlines():
        if line.startswith("# "):
            headings.append((1, line[2:].strip()))
        elif line.startswith("## "):
            headings.append((2, line[3:].strip()))
    return headings


def _normalize_for_page_lookup(text: str) -> str:
    normalized = re.sub(r"\s+", " ", text.replace("\u00a0", " ")).strip().lower()
    return normalized


def _rendered_page_heading_candidates(raw_text: str) -> set[str]:
    lines = [_normalize_for_page_lookup(line) for line in raw_text.splitlines() if line.strip()]
    candidates = set(lines)
    for start in range(len(lines)):
        combined = lines[start]
        for end in range(start + 1, min(start + 4, len(lines))):
            combined = f"{combined} {lines[end]}"
            candidates.add(combined)
    return candidates


def _rendered_page_has_exact_heading(raw_text: str, heading: str) -> bool:
    needle = _normalize_for_page_lookup(heading)
    return needle in _rendered_page_heading_candidates(raw_text)


def _rendered_page_looks_like_contents(raw_text: str, headings: list[tuple[int, str]]) -> bool:
    lines = [_normalize_for_page_lookup(line) for line in raw_text.splitlines() if line.strip()]
    if any(line == "contents" for line in lines):
        return True
    dot_leader_lines = sum(1 for line in raw_text.splitlines() if re.search(r"\.{4,}\s*\d+\s*$", line))
    return dot_leader_lines >= 3


def _rendered_body_start_page(page_texts: list[str], first_heading: str, headings: list[tuple[int, str]]) -> int:
    for page_index, raw_text in enumerate(page_texts, start=1):
        if _rendered_page_has_exact_heading(raw_text, first_heading) and not _rendered_page_looks_like_contents(raw_text, headings):
            return page_index
    return 3


def _rendered_title_block_page_number(raw_text: str, physical_page_index: int) -> str:
    lines = [line.rstrip() for line in raw_text.splitlines()]
    for line_index, line in enumerate(lines):
        if not re.search(r"\bPage\b", line):
            continue
        for following in lines[line_index + 1 : line_index + 6]:
            matches = re.findall(r"\b(\d{1,3})\s*$", following)
            if not matches:
                continue
            value = int(matches[-1])
            if value >= 3:
                return str(value)
    return str(physical_page_index)


def _applescript_string(value: Path | str) -> str:
    return str(value).replace("\\", "\\\\").replace('"', '\\"')


def _render_pdf_for_toc_with_word(docx_path: Path, output_dir: Path) -> Path | None:
    osascript = shutil.which("osascript")
    if not osascript or not Path("/Applications/Microsoft Word.app").exists():
        return None
    output_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = output_dir / f"{docx_path.stem}.pdf"
    pdf_path.unlink(missing_ok=True)
    script = f'''
set docPath to POSIX file "{_applescript_string(docx_path)}"
set outPath to "{_applescript_string(pdf_path)}"
tell application "Microsoft Word"
    set display alerts to none
    open docPath
    delay 1
    set docRef to active document
    save as docRef file name outPath file format format PDF
    close docRef saving no
end tell
'''
    subprocess.run(
        [osascript],
        input=script,
        check=False,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        timeout=180,
    )
    return pdf_path if pdf_path.exists() else None


def _render_pdf_for_toc(docx_path: Path, output_dir: Path) -> Path | None:
    word_pdf = _render_pdf_for_toc_with_word(docx_path, output_dir)
    if word_pdf is not None:
        return word_pdf
    soffice = shutil.which("libreoffice") or shutil.which("soffice")
    if not soffice:
        return None
    output_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = output_dir / f"{docx_path.stem}.pdf"
    pdf_path.unlink(missing_ok=True)
    result = subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(output_dir), str(docx_path)],
        check=False,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )
    if result.returncode != 0 or not pdf_path.exists():
        return None
    return pdf_path


def _pdf_pages_text_for_toc(pdf_path: Path) -> list[str]:
    pdftotext = shutil.which("pdftotext")
    if not pdftotext or not pdf_path.exists():
        return []
    result = subprocess.run(
        [pdftotext, "-layout", str(pdf_path), "-"],
        check=False,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        timeout=120,
    )
    if result.returncode != 0:
        return []
    pages = result.stdout.split("\f")
    if pages and pages[-1] == "":
        pages.pop()
    return pages


def _extract_rendered_heading_pages(docx_path: Path, headings: list[tuple[int, str]]) -> dict[str, str]:
    pdf_path = _render_pdf_for_toc(docx_path, DRAFT_DIR / "toc_probe")
    if pdf_path is None:
        return {}
    raw_page_texts = _pdf_pages_text_for_toc(pdf_path)
    if not raw_page_texts:
        return {}
    body_start_page = _rendered_body_start_page(raw_page_texts, headings[0][1], headings) if headings else 3
    page_heading_candidates = [_rendered_page_heading_candidates(raw_text) for raw_text in raw_page_texts]
    heading_pages: dict[str, str] = {}
    for _level, heading in headings:
        needle = _normalize_for_page_lookup(heading)
        if not needle:
            continue
        for page_index, heading_candidates in enumerate(page_heading_candidates[body_start_page - 1 :], start=body_start_page):
            if needle in heading_candidates:
                heading_pages[heading] = _rendered_title_block_page_number(
                    raw_page_texts[page_index - 1],
                    page_index,
                )
                break
    return heading_pages


def _add_contents_entry(doc: Document, title: str, page: str) -> None:
    paragraph = doc.add_paragraph(style="Contents Entry")
    paragraph.paragraph_format.first_line_indent = Mm(0)
    paragraph.paragraph_format.left_indent = Mm(0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Mm(165), alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.DOTS
    )
    run = paragraph.add_run(title)
    _set_run_font(run, 13)
    paragraph.add_run("\t")
    page_run = paragraph.add_run(page)
    _set_run_font(page_run, 13)


def _add_contents_pages(doc: Document, headings: list[tuple[int, str]], contents_pages: dict[str, str]) -> bool:
    overflow = len(headings) > CONTENTS_FIRST_PAGE_ENTRY_LIMIT
    contents = doc.add_paragraph(style="Contents Title")
    contents.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = contents.add_run("Contents")
    _set_run_font(title_run, 14, bold=True)
    doc.add_paragraph()
    for index, (_level, heading) in enumerate(headings):
        if index == CONTENTS_FIRST_PAGE_ENTRY_LIMIT:
            doc.add_section(WD_SECTION.NEW_PAGE)
        _add_contents_entry(doc, heading, contents_pages.get(heading, "3"))
    return overflow


def _add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Body Text")
    paragraph.paragraph_format.first_line_indent = Mm(12.5)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(_keep_numbered_reference_together(text))
    _set_run_font(run, 13)


def _add_float_spacing_paragraph(doc: Document) -> None:
    paragraph = doc.add_paragraph(style="Body Text")
    paragraph.paragraph_format.first_line_indent = Mm(0)
    paragraph.paragraph_format.left_indent = Mm(0)
    paragraph.paragraph_format.right_indent = Mm(0)
    paragraph.paragraph_format.line_spacing = Pt(13)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def _add_references_entry(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="References Entry")
    paragraph.paragraph_format.first_line_indent = Mm(0)
    paragraph.paragraph_format.left_indent = Mm(0)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    _set_run_font(run, 13)


def _add_formula_explanation_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Body Text")
    paragraph.paragraph_format.first_line_indent = Mm(0)
    paragraph.paragraph_format.left_indent = Mm(0)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Mm(FORMULA_SYMBOL_TAB_MM), alignment=WD_TAB_ALIGNMENT.LEFT, leader=WD_TAB_LEADER.SPACES
    )
    if text.startswith("where "):
        where_run = paragraph.add_run("where ")
        _set_run_font(where_run, 13)
        paragraph.add_run("\t")
        symbol_text = text[len("where ") :]
    else:
        paragraph.add_run("\t")
        symbol_text = text
    run = paragraph.add_run(_keep_numbered_reference_together(symbol_text))
    _set_run_font(run, 13)


def _add_table_caption(doc: Document, text: str) -> None:
    if text.startswith(("Table 1.1", "Table 5.1")):
        doc.add_page_break()
    else:
        _add_float_spacing_paragraph(doc)
    paragraph = doc.add_paragraph(style="Table Caption")
    paragraph.paragraph_format.first_line_indent = Mm(0)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(_keep_numbered_reference_together(text))
    _set_run_font(run, 13)


def _add_caption(doc: Document, text: str, style: str) -> None:
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.first_line_indent = Mm(0)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if style == "Figure Caption" else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(_keep_numbered_reference_together(text))
    _set_run_font(run, 13)


def _add_figure(doc: Document, figure_id: str) -> None:
    path, caption, width_mm = FIGURE_DEFINITIONS[figure_id]
    if not path.exists():
        raise FileNotFoundError(f"Missing {figure_id} image: {path}")
    _add_float_spacing_paragraph(doc)
    image_paragraph = doc.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.first_line_indent = Mm(0)
    image_paragraph.paragraph_format.space_before = Pt(0)
    image_paragraph.paragraph_format.space_after = Pt(0)
    image_paragraph.paragraph_format.keep_with_next = True
    image_paragraph.add_run().add_picture(str(path), width=Mm(width_mm))
    _add_caption(doc, caption, "Figure Caption")
    _add_float_spacing_paragraph(doc)


def _add_formula_paragraph(doc: Document, equation_omml: str, number: str, trailing_punctuation: str = "") -> None:
    paragraph = doc.add_paragraph(style="Formula")
    paragraph.paragraph_format.first_line_indent = Mm(0)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(13)
    paragraph.paragraph_format.space_after = Pt(13)
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Mm(82.5), alignment=WD_TAB_ALIGNMENT.CENTER, leader=WD_TAB_LEADER.SPACES
    )
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Mm(165), alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.SPACES
    )
    paragraph.add_run("\t")
    paragraph._p.append(parse_xml(equation_omml))
    if trailing_punctuation:
        punctuation_run = paragraph.add_run(trailing_punctuation)
        _set_run_font(punctuation_run, 13)
    paragraph.add_run("\t")
    number_run = paragraph.add_run(number)
    _set_run_font(number_run, 13)


def _formula_omml(text: str) -> str:
    return (
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        f"<m:r><m:t>{text}</m:t></m:r>"
        "</m:oMath>"
    )


def _m_text(text: str) -> str:
    space = ' xml:space="preserve"' if text.startswith(" ") or text.endswith(" ") else ""
    return f"<m:r><m:t{space}>{text}</m:t></m:r>"


def _m_sub(base: str, subscript: str) -> str:
    return (
        "<m:sSub>"
        f"<m:e>{_m_text(base)}</m:e>"
        f"<m:sub>{_m_text(subscript)}</m:sub>"
        "</m:sSub>"
    )


def _m_sup(base: str, superscript: str) -> str:
    return (
        "<m:sSup>"
        f"<m:e>{_m_text(base)}</m:e>"
        f"<m:sup>{_m_text(superscript)}</m:sup>"
        "</m:sSup>"
    )


def _m_subsup(base: str, subscript: str, superscript: str) -> str:
    return (
        "<m:sSubSup>"
        f"<m:e>{_m_text(base)}</m:e>"
        f"<m:sub>{_m_text(subscript)}</m:sub>"
        f"<m:sup>{_m_text(superscript)}</m:sup>"
        "</m:sSubSup>"
    )


def _m_frac(numerator: str, denominator: str) -> str:
    return (
        "<m:f>"
        '<m:fPr><m:type m:val="bar"/></m:fPr>'
        f"<m:num>{numerator}</m:num>"
        f"<m:den>{denominator}</m:den>"
        "</m:f>"
    )


def _m_omath(*parts: str) -> str:
    return '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">' + "".join(parts) + "</m:oMath>"


def _formula_4_1_omml() -> str:
    return _m_omath(
        _m_sub("I", "pullup"),
        _m_text("="),
        _m_frac(_m_sub("U", "DD"), _m_sub("R", "pullup")),
    )


def _formula_4_2_omml() -> str:
    return _m_omath(
        _m_sub("U", "avg"),
        _m_text("="),
        _m_text("D"),
        _m_sub("U", "s"),
    )


def _formula_4_3_omml() -> str:
    return _m_omath(
        _m_sub("P", "load"),
        _m_text("="),
        _m_frac(_m_subsup("U", "s", "2"), _m_sub("R", "load")),
    )


def _formula_4_4_omml() -> str:
    return _m_omath(
        _m_sub("P", "Q"),
        _m_text("="),
        _m_subsup("I", "load", "2"),
        _m_sub("R", "DS(on)"),
    )


def _formula_4_5_omml() -> str:
    return _m_omath(
        _m_text("e(t)="),
        _m_sub("T", "set"),
        _m_text("−T(t)"),
    )


def _formula_pid_omml() -> str:
    return (
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        '<m:r><m:t>u(t)=</m:t></m:r>'
        "<m:sSub>"
        "<m:e><m:r><m:t>K</m:t></m:r></m:e>"
        "<m:sub><m:r><m:t>p</m:t></m:r></m:sub>"
        "</m:sSub>"
        '<m:r><m:t xml:space="preserve"> e(t)+</m:t></m:r>'
        "<m:sSub>"
        "<m:e><m:r><m:t>K</m:t></m:r></m:e>"
        "<m:sub><m:r><m:t>i</m:t></m:r></m:sub>"
        "</m:sSub>"
        '<m:r><m:t xml:space="preserve"> </m:t></m:r>'
        "<m:nary>"
        "<m:naryPr>"
        '<m:chr m:val="∫"/>'
        '<m:limLoc m:val="subSup"/>'
        "</m:naryPr>"
        "<m:sub><m:r><m:t>0</m:t></m:r></m:sub>"
        "<m:sup><m:r><m:t>t</m:t></m:r></m:sup>"
        '<m:e><m:r><m:t xml:space="preserve">e(τ) dτ</m:t></m:r></m:e>'
        "</m:nary>"
        '<m:r><m:t xml:space="preserve">+</m:t></m:r>'
        "<m:sSub>"
        "<m:e><m:r><m:t>K</m:t></m:r></m:e>"
        "<m:sub><m:r><m:t>d</m:t></m:r></m:sub>"
        "</m:sSub>"
        '<m:r><m:t xml:space="preserve"> </m:t></m:r>'
        "<m:f>"
        "<m:fPr><m:type m:val=\"bar\"/></m:fPr>"
        '<m:num><m:r><m:t xml:space="preserve">d e(t)</m:t></m:r></m:num>'
        "<m:den><m:r><m:t>dt</m:t></m:r></m:den>"
        "</m:f>"
        "</m:oMath>"
    )


def _formula_ziegler_nichols_omml() -> str:
    return _m_omath(
        _m_sub("K", "p"),
        _m_text("=0.6"),
        _m_sub("K", "u"),
        _m_text("; "),
        _m_sub("K", "i"),
        _m_text("="),
        _m_frac(_m_text("1.2") + _m_sub("K", "u"), _m_sub("T", "u")),
        _m_text("; "),
        _m_sub("K", "d"),
        _m_text("=0.075"),
        _m_sub("K", "u"),
        _m_sub("T", "u"),
    )


def _formula_4_8_omml() -> str:
    return _m_omath(
        _m_text("D="),
        _m_frac(_m_text("u"), _m_text("255")),
    )


def _add_formula_by_id(doc: Document, formula_id: str) -> None:
    formulas = {
        "FORMULA_3_1": (_formula_4_1_omml(), "(3.1)"),
        "FORMULA_3_2": (_formula_4_2_omml(), "(3.2)"),
        "FORMULA_3_3": (_formula_4_3_omml(), "(3.3)"),
        "FORMULA_3_4": (_formula_4_4_omml(), "(3.4)"),
        "FORMULA_3_5": (_formula_4_5_omml(), "(3.5)"),
        "FORMULA_3_6": (_formula_pid_omml(), "(3.6)"),
        "FORMULA_3_7": (_formula_ziegler_nichols_omml(), "(3.7)"),
        "FORMULA_3_8": (_formula_4_8_omml(), "(3.8)"),
    }
    equation, number = formulas[formula_id]
    _add_formula_paragraph(doc, equation, number, trailing_punctuation=",")


def _split_markdown_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _is_markdown_table_separator(line: str) -> bool:
    cells = _split_markdown_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def _set_cell_width(cell, width_mm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(docx_qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(docx_qn("w:type"), "dxa")
    tc_w.set(docx_qn("w:w"), str(int(round(width_mm * 1440 / 25.4))))


def _format_table_cell(cell, text: str, *, bold: bool = False, width_mm: float | None = None) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    if width_mm is not None:
        cell.width = Mm(width_mm)
        _set_cell_width(cell, width_mm)
    paragraph = cell.paragraphs[0]
    paragraph.style = "Normal"
    paragraph.paragraph_format.first_line_indent = Mm(0)
    paragraph.paragraph_format.left_indent = Mm(0)
    paragraph.paragraph_format.right_indent = Mm(0)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.clear()
    run = paragraph.add_run(text.strip() or "-")
    _set_run_font(run, 13, bold=bold)


def _set_table_width(table, width_mm: float) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(docx_qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(docx_qn("w:type"), "dxa")
    tbl_w.set(docx_qn("w:w"), str(int(round(width_mm * 1440 / 25.4))))
    tbl_ind = tbl_pr.find(docx_qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(docx_qn("w:type"), "dxa")
    tbl_ind.set(docx_qn("w:w"), "0")


def _set_table_row_pagination(row, *, header: bool = False) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(docx_qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))
    if header and tr_pr.find(docx_qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def _add_table_continuation_label(doc: Document, table_number: str, *, page_break_before: bool = False) -> None:
    paragraph = doc.add_paragraph(style="Body Text")
    paragraph.paragraph_format.first_line_indent = Mm(0)
    paragraph.paragraph_format.left_indent = Mm(0)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.page_break_before = page_break_before
    run = paragraph.add_run(_keep_numbered_reference_together(f"Continue of the Table {table_number}"))
    _set_run_font(run, 13)


def _add_word_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    column_widths = [38.0, 87.0, 40.0]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _set_table_width(table, sum(column_widths[: len(headers)]))
    _set_table_row_pagination(table.rows[0], header=True)
    for idx, header in enumerate(headers):
        width = column_widths[idx] if idx < len(column_widths) else 165.0 / len(headers)
        _format_table_cell(table.rows[0].cells[idx], _keep_numbered_reference_together(header), bold=True, width_mm=width)
    for row_values in rows:
        row = table.add_row()
        _set_table_row_pagination(row)
        row_cells = row.cells
        for idx, cell in enumerate(row_cells):
            value = row_values[idx] if idx < len(row_values) else "-"
            width = column_widths[idx] if idx < len(column_widths) else 165.0 / len(headers)
            _format_table_cell(cell, _keep_numbered_reference_together(value), width_mm=width)


def _add_markdown_table(doc: Document, lines: list[str], caption_text: str | None = None) -> None:
    if len(lines) < 3:
        for line in lines:
            _add_body_paragraph(doc, line)
        return
    headers = _split_markdown_table_row(lines[0])
    rows = [_split_markdown_table_row(line) for line in lines[2:]]
    split_after = None
    table_number = None
    if caption_text:
        table_match = re.match(r"^(Table\s+\d+\.\d+)", caption_text)
        if table_match:
            table_id = table_match.group(1)
            split_after = MANUAL_TABLE_SPLITS.get(table_id)
            table_number = table_id.split()[1]
    if split_after is not None and 0 < split_after < len(rows) and table_number:
        _add_word_table(doc, headers, rows[:split_after])
        _add_table_continuation_label(doc, table_number, page_break_before=True)
        _add_word_table(doc, headers, rows[split_after:])
        return
    _add_word_table(doc, headers, rows)


def _add_list_item(doc: Document, text: str) -> None:
    clean = text.strip()
    if not clean.endswith((".", ";", ":")):
        clean += ";"
    _add_body_paragraph(doc, f"- {clean}")


def _add_heading(doc: Document, level: int, text: str) -> None:
    style = "Heading 1" if level == 1 else "Heading 2"
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.first_line_indent = Mm(12.5)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.keep_with_next = True
    if level == 1:
        paragraph.paragraph_format.page_break_before = True
        paragraph.paragraph_format.space_after = Pt(13)
        run = paragraph.add_run(text.upper())
        _set_run_font(run, 14, bold=True)
    else:
        paragraph.paragraph_format.space_before = Pt(13)
        paragraph.paragraph_format.space_after = Pt(13)
        run = paragraph.add_run(text)
        _set_run_font(run, 13, bold=True)


def _add_markdown_body(doc: Document, markdown: str) -> None:
    pending_paragraph: list[str] = []
    inserted_figures: set[str] = set()
    after_formula = False
    in_formula_explanation = False
    last_table_caption: str | None = None

    def flush() -> None:
        nonlocal after_formula, in_formula_explanation
        if pending_paragraph:
            paragraph_text = " ".join(pending_paragraph).strip()
            if after_formula and paragraph_text.startswith("where "):
                _add_formula_explanation_paragraph(doc, paragraph_text)
                in_formula_explanation = True
            elif in_formula_explanation and re.match(r"^[A-Za-z][A-Za-z0-9_]*(?:\([^)]*\))?\s+\u2013\s+", paragraph_text):
                _add_formula_explanation_paragraph(doc, paragraph_text)
            else:
                _add_body_paragraph(doc, paragraph_text)
                in_formula_explanation = False
            after_formula = False
            for figure_id in FIGURE_DEFINITIONS:
                if figure_id in paragraph_text and figure_id not in inserted_figures:
                    _add_figure(doc, figure_id)
                    inserted_figures.add(figure_id)
            pending_paragraph.clear()

    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            flush()
            index += 1
            continue
        if line.startswith("# "):
            flush()
            _add_heading(doc, 1, line[2:].strip())
            index += 1
            continue
        if line.startswith("## "):
            flush()
            _add_heading(doc, 2, line[3:].strip())
            index += 1
            continue
        if re.match(r"^Table\s+\d+(?:\.\d+)?\s+–\s+\S", line):
            flush()
            _add_table_caption(doc, line)
            last_table_caption = line
            index += 1
            continue
        if (
            line.startswith("|")
            and index + 1 < len(lines)
            and _is_markdown_table_separator(lines[index + 1].strip())
        ):
            flush()
            table_lines = [line, lines[index + 1].strip()]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            _add_markdown_table(doc, table_lines, last_table_caption)
            _add_float_spacing_paragraph(doc)
            last_table_caption = None
            continue
        if line.startswith("- "):
            flush()
            _add_list_item(doc, line[2:].strip())
            index += 1
            continue
        if re.match(r"^\[\d+\]\s+\S", line):
            flush()
            _add_references_entry(doc, line)
            index += 1
            continue
        if line.startswith("{{FORMULA_") and line.endswith("}}"):
            flush()
            _add_formula_by_id(doc, line.strip("{}"))
            after_formula = True
            in_formula_explanation = False
            index += 1
            continue
        pending_paragraph.append(line)
        index += 1
    flush()


def build_raw_docx(path: Path = RAW_DOCX, contents_pages: dict[str, str] | None = None) -> None:
    markdown = _read_markdown_source()
    headings = _extract_headings(markdown)
    contents_pages = contents_pages or {}

    doc = Document()
    _set_doc_layout(doc)
    _configure_styles(doc)

    # Cover section: intentionally blank; postprocess supplies the accepted border-only frame.
    doc.add_paragraph()
    doc.add_section(WD_SECTION.NEW_PAGE)

    contents_overflow = _add_contents_pages(doc, headings, contents_pages)

    if not contents_overflow:
        doc.add_section(WD_SECTION.NEW_PAGE)
    _add_markdown_body(doc, markdown)

    _set_doc_layout(doc)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def build_draft(*, partial: bool = True) -> int:
    if SAMPLE_TEMPLATE.exists():
        DRAFT_DIR.mkdir(parents=True, exist_ok=True)
        shutil.copy2(SAMPLE_TEMPLATE, FINAL_DOCX)
    build_raw_docx(RAW_DOCX)
    postprocess_docx(RAW_DOCX, FINAL_DOCX, TEMPLATE_DIR)
    markdown = _read_markdown_source()
    heading_pages = _extract_rendered_heading_pages(FINAL_DOCX, _extract_headings(markdown))
    if heading_pages:
        build_raw_docx(RAW_DOCX, heading_pages)
        postprocess_docx(RAW_DOCX, FINAL_DOCX, TEMPLATE_DIR)
    else:
        print("WARNING: Could not resolve rendered Contents page numbers; kept placeholder page numbers.")
    reporter = check_docx(FINAL_DOCX, REPORT_DOCX, partial=partial)
    counts = reporter.counts()
    print(f"Wrote {RAW_DOCX}")
    print(f"Wrote {FINAL_DOCX}")
    print(f"Wrote {REPORT_DOCX}")
    print(f"PASS: {counts['PASS']}")
    print(f"WARNING: {counts['WARNING']}")
    print(f"ERROR: {counts['ERROR']}")
    return 1 if reporter.has_errors else 0


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--complete", action="store_true", help="Run complete-thesis checks instead of partial-draft checks")
    args = parser.parse_args(argv)
    return build_draft(partial=not args.complete)


if __name__ == "__main__":
    raise SystemExit(main())
